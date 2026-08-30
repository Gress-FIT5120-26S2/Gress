from docx import Document


document = Document(r"docs/photo-recognition/食物新鲜度识别模型技术使用文档.docx")

for paragraph in document.paragraphs:
    if paragraph.text.strip():
        print(paragraph.text)

for table_index, table in enumerate(document.tables, start=1):
    print(f"\nTABLE {table_index}")
    for row in table.rows:
        print(" | ".join(cell.text.replace("\n", " / ") for cell in row.cells))
